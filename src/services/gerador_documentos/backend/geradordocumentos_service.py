from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile
import threading
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

import pandas as pd
from docx import Document


MODELO_WORD_NOME = "MODELO JUNTADA.docx"
PLANILHA_EXTENSOES = {".xlsx", ".xlsm", ".xls", ".csv"}

COLUNAS_OBRIGATORIAS = [
    "OJ NUMERO",
    "OJ DESCRICAO",
    "CIDADE",
    "UF",
    "N\u00daMERO DO PROCESSO",
    "AUTOR",
]

OAB_POR_UF = {
    "SP": "OAB/SP n\u00ba 403.110",
    "AC": "OAB/AC n\u00ba 6205",
    "AL": "OAB/AL n\u00ba 19.554A",
    "AM": "OAB/AM n\u00ba A1814",
    "AP": "OAB/AP n\u00ba 5352-A",
    "BA": "OAB/BA n\u00ba 72947",
    "CE": "OAB/CE n\u00ba 47894-A",
    "DF": "OAB/DF n\u00ba 71599",
    "ES": "OAB/ES n\u00ba 37035",
    "GO": "OAB/GO n\u00ba 65009-A",
    "MA": "OAB/MA n\u00ba 25106-A",
    "MG": "OAB/MG n\u00ba 217298",
    "MS": "OAB/MS n\u00ba 27556\u00aa",
    "MT": "OAB/MT n\u00ba 32134-A",
    "PA": "OAB/PA n\u00ba 34787-A",
    "PB": "OAB/PB n\u00ba 30904-A",
    "PE": "OAB/PE n\u00ba 58145",
    "PI": "OAB/PI n\u00ba 21463",
    "PR": "OAB/PR n\u00ba 109831",
    "RJ": "OAB/RJ n\u00ba 233392",
    "RN": "OAB/RN n\u00ba 20.111-A",
    "RO": "OAB/RO n\u00ba 12423",
    "RR": "OAB/RR n\u00ba 694-A",
    "RS": "OAB/RS n\u00ba 127170A",
    "SC": "OAB/SC n\u00ba 65177",
    "SE": "OAB/SE n\u00ba 1485A",
    "TO": "OAB/TO n\u00ba 11.641-A",
}

StatusCallback = Callable[[str, int], None]

jobs: dict[str, dict[str, Any]] = {}
jobs_lock = threading.Lock()


def _runtime_base() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[4]


def _resource_candidates(nome_arquivo: str) -> list[Path]:
    candidates: list[Path] = []

    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        base = Path(meipass)
        candidates.extend([
            base / "src" / "docs" / nome_arquivo,
            base / "docs" / nome_arquivo,
            base / nome_arquivo,
        ])

    runtime_base = _runtime_base()
    here = Path(__file__).resolve().parent
    candidates.extend([
        runtime_base / "src" / "docs" / nome_arquivo,
        runtime_base / "docs" / nome_arquivo,
        here / nome_arquivo,
    ])

    unique: list[Path] = []
    seen: set[str] = set()
    for candidate in candidates:
        key = str(candidate)
        if key not in seen:
            unique.append(candidate)
            seen.add(key)

    return unique


def caminho_recurso(nome_arquivo: str) -> Path:
    candidates = _resource_candidates(nome_arquivo)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def caminho_modelo_word() -> Path:
    return caminho_recurso(MODELO_WORD_NOME)


def modelo_info() -> dict[str, Any]:
    modelo = caminho_modelo_word()
    return {
        "name": modelo.name,
        "path": str(modelo),
        "exists": modelo.exists(),
        "colunas_obrigatorias": list(COLUNAS_OBRIGATORIAS),
    }


def _path_obrigatorio(valor: Any, nome: str) -> Path:
    texto = str(valor or "").strip()
    if not texto:
        raise ValueError(f"{nome} nao informado.")
    return Path(texto).expanduser().resolve()


def validar_pdf_disponivel() -> None:
    if os.name != "nt":
        raise RuntimeError("A geracao em PDF usa Microsoft Word via COM e deve ser executada no Windows.")

    try:
        import pythoncom  # noqa: F401
        import win32com.client  # noqa: F401
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Dependencia ausente para conversao em PDF. Instale 'pywin32' e gere o executavel novamente."
        ) from exc


def converter_docx_para_pdf(caminho_docx: Path, caminho_pdf: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Dependencia ausente para conversao em PDF. Instale 'pywin32' e gere o executavel novamente."
        ) from exc

    wd_format_pdf = 17
    word = None
    doc = None

    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(str(caminho_docx))
        doc.SaveAs(str(caminho_pdf), FileFormat=wd_format_pdf)
    except Exception as exc:
        raise RuntimeError(f"Falha ao converter o arquivo para PDF: {caminho_docx.name}\n{exc}") from exc
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass

        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass

        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def normalizar_valor(valor: Any) -> str:
    if pd.isna(valor):
        return ""

    if isinstance(valor, pd.Timestamp):
        return valor.strftime("%d/%m/%Y")

    return str(valor).strip()


def ler_planilha(caminho_planilha: Path) -> pd.DataFrame:
    sufixo = caminho_planilha.suffix.lower()

    if sufixo in {".xlsx", ".xlsm", ".xls"}:
        df = pd.read_excel(caminho_planilha)
    elif sufixo == ".csv":
        try:
            df = pd.read_csv(caminho_planilha, sep=";", dtype=str)
        except Exception:
            df = pd.read_csv(caminho_planilha, sep=",", dtype=str)
    else:
        raise ValueError(
            f"Formato de planilha nao suportado: {caminho_planilha.suffix}. Use .xlsx, .xls, .xlsm ou .csv"
        )

    df.columns = [str(col).strip() for col in df.columns]

    colunas_faltando = [col for col in COLUNAS_OBRIGATORIAS if col not in df.columns]
    if colunas_faltando:
        raise ValueError(
            "A planilha nao contem todas as colunas obrigatorias.\n"
            f"Faltando: {', '.join(colunas_faltando)}"
        )

    return df


def montar_substituicoes(linha: pd.Series) -> dict[str, str]:
    substituicoes: dict[str, str] = {}

    for coluna in COLUNAS_OBRIGATORIAS:
        valor = normalizar_valor(linha.get(coluna, ""))

        if coluna in {"CIDADE", "UF"}:
            valor = valor.upper()

        substituicoes[coluna] = valor

    uf = substituicoes.get("UF", "")
    substituicoes["NUMERO"] = OAB_POR_UF.get(uf, "")

    return substituicoes


def substituir_alvo_em_paragrafo(paragrafo, alvo: str, valor: str) -> None:
    if not alvo:
        return

    if not paragrafo.runs:
        if paragrafo.text:
            paragrafo.text = paragrafo.text.replace(alvo, valor)
        return

    for run in paragrafo.runs:
        if alvo in run.text:
            run.text = run.text.replace(alvo, valor)

    while True:
        texto_completo = "".join(run.text for run in paragrafo.runs)
        inicio = texto_completo.find(alvo)

        if inicio == -1:
            break

        fim = inicio + len(alvo)
        mapa_caracteres = []

        for indice_run, run in enumerate(paragrafo.runs):
            for indice_char, _ in enumerate(run.text):
                mapa_caracteres.append((indice_run, indice_char))

        if fim > len(mapa_caracteres):
            break

        primeiro_run, primeiro_char = mapa_caracteres[inicio]
        ultimo_run, ultimo_char = mapa_caracteres[fim - 1]

        prefixo = paragrafo.runs[primeiro_run].text[:primeiro_char]
        sufixo = paragrafo.runs[ultimo_run].text[ultimo_char + 1:]

        if primeiro_run == ultimo_run:
            paragrafo.runs[primeiro_run].text = prefixo + valor + sufixo
        else:
            paragrafo.runs[primeiro_run].text = prefixo + valor

            for i in range(primeiro_run + 1, ultimo_run):
                paragrafo.runs[i].text = ""

            paragrafo.runs[ultimo_run].text = sufixo


def substituir_em_paragrafo(paragrafo, substituicoes: dict[str, str]) -> None:
    if not paragrafo.text and not paragrafo.runs:
        return

    for chave, valor in substituicoes.items():
        alvos = [
            f"{{{{{chave}}}}}",
            f"{{{chave}}}",
            chave,
        ]

        for alvo in alvos:
            substituir_alvo_em_paragrafo(paragrafo, alvo, valor)


def substituir_em_tabela(tabela, substituicoes: dict[str, str]) -> None:
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                substituir_em_paragrafo(paragrafo, substituicoes)

            for tabela_interna in celula.tables:
                substituir_em_tabela(tabela_interna, substituicoes)


def substituir_em_container(container, substituicoes: dict[str, str]) -> None:
    for paragrafo in container.paragraphs:
        substituir_em_paragrafo(paragrafo, substituicoes)

    for tabela in container.tables:
        substituir_em_tabela(tabela, substituicoes)


def substituir_no_documento(doc: Document, substituicoes: dict[str, str]) -> None:
    substituir_em_container(doc, substituicoes)

    for secao in doc.sections:
        substituir_em_container(secao.header, substituicoes)
        substituir_em_container(secao.footer, substituicoes)


def limpar_nome_arquivo(nome: str) -> str:
    nome = re.sub(r'[<>:"/\\|?*]', "_", nome)
    nome = re.sub(r"\s+", " ", nome).strip()
    return nome[:180] if nome else "documento"


def montar_nome_base(indice: int, linha: pd.Series) -> str:
    processo = normalizar_valor(linha.get("N\u00daMERO DO PROCESSO", ""))
    nome_base = limpar_nome_arquivo(processo)
    return nome_base or f"documento_{indice:04d}"


def caminho_unico(caminho: Path) -> Path:
    if not caminho.exists():
        return caminho

    contador = 1
    while True:
        candidato = caminho.with_name(f"{caminho.stem} ({contador}){caminho.suffix}")
        if not candidato.exists():
            return candidato
        contador += 1


def criar_pasta_lote(pasta_base: Path) -> Path:
    timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    pasta_lote = pasta_base / f"Documentos Gerados - {timestamp}"
    pasta_lote.mkdir(parents=True, exist_ok=True)
    return pasta_lote


def validar_entrada(caminho_planilha: Path, pasta_saida_base: Path, gerar_pdf: bool) -> dict[str, Any]:
    if not caminho_planilha.exists() or not caminho_planilha.is_file():
        raise FileNotFoundError(f"Planilha nao encontrada: {caminho_planilha}")

    if caminho_planilha.suffix.lower() not in PLANILHA_EXTENSOES:
        raise ValueError("Selecione uma planilha .xlsx, .xls, .xlsm ou .csv.")

    modelo = caminho_modelo_word()
    if not modelo.exists():
        candidatos = "\n".join(str(p) for p in _resource_candidates(MODELO_WORD_NOME))
        raise FileNotFoundError(
            f"Modelo Word fixo nao encontrado: {MODELO_WORD_NOME}\nCaminhos verificados:\n{candidatos}"
        )

    pasta_saida_base.mkdir(parents=True, exist_ok=True)
    if not pasta_saida_base.is_dir():
        raise NotADirectoryError(f"Pasta de saida invalida: {pasta_saida_base}")

    df = ler_planilha(caminho_planilha)

    if gerar_pdf:
        validar_pdf_disponivel()

    return {
        "linhas": int(len(df)),
        "modelo": str(modelo),
        "saida": str(pasta_saida_base),
        "gerar_pdf": bool(gerar_pdf),
    }


def validar_payload(planilha: Any, pasta_saida: Any, gerar_pdf: bool) -> dict[str, Any]:
    caminho_planilha = _path_obrigatorio(planilha, "Planilha")
    pasta_saida_base = _path_obrigatorio(pasta_saida, "Pasta de saida")
    return validar_entrada(caminho_planilha, pasta_saida_base, bool(gerar_pdf))


def gerar_documentos(
    caminho_planilha: Path,
    pasta_saida_base: Path,
    gerar_pdf: bool = False,
    callback_status: StatusCallback | None = None,
    cancel_event: threading.Event | None = None,
) -> tuple[int, Path, bool]:
    validar_entrada(caminho_planilha, pasta_saida_base, gerar_pdf)

    pasta_saida = criar_pasta_lote(pasta_saida_base)
    df = ler_planilha(caminho_planilha)
    modelo = caminho_modelo_word()

    total = 0
    cancelado = False

    for indice, (_, linha) in enumerate(df.iterrows(), start=1):
        if cancel_event and cancel_event.is_set():
            cancelado = True
            break

        substituicoes = montar_substituicoes(linha)

        doc = Document(modelo)
        substituir_no_documento(doc, substituicoes)

        nome_base = montar_nome_base(indice, linha)
        caminho_docx = caminho_unico(pasta_saida / f"{nome_base}.docx")
        doc.save(caminho_docx)

        arquivos = [caminho_docx.name]
        if gerar_pdf:
            caminho_pdf = caminho_unico(caminho_docx.with_suffix(".pdf"))
            converter_docx_para_pdf(caminho_docx, caminho_pdf)
            arquivos.append(caminho_pdf.name)

        total += 1

        if callback_status:
            callback_status(f"Gerado: {', '.join(arquivos)}", total)

    return total, pasta_saida, cancelado


def _serializar_job(job: dict[str, Any]) -> dict[str, Any]:
    retorno = {k: v for k, v in job.items() if k not in {"cancel_event", "thread"}}
    retorno["logs"] = list(retorno.get("logs") or [])
    return retorno


def iniciar_job(planilha: Any, pasta_saida: Any, gerar_pdf: bool = False) -> dict[str, Any]:
    caminho_planilha = _path_obrigatorio(planilha, "Planilha")
    pasta_saida_base = _path_obrigatorio(pasta_saida, "Pasta de saida")

    info = validar_entrada(caminho_planilha, pasta_saida_base, bool(gerar_pdf))
    job_id = str(uuid.uuid4())
    cancel_event = threading.Event()

    job: dict[str, Any] = {
        "id": job_id,
        "status": "processando",
        "planilha": str(caminho_planilha),
        "pasta_saida_base": str(pasta_saida_base),
        "pasta_final": "",
        "gerar_pdf": bool(gerar_pdf),
        "total": 0,
        "linhas": info["linhas"],
        "cancelado": False,
        "erro": "",
        "logs": ["Processando documentos..."],
        "cancel_event": cancel_event,
    }

    with jobs_lock:
        jobs[job_id] = job

    def callback_status(mensagem: str, total: int) -> None:
        with jobs_lock:
            current = jobs.get(job_id)
            if not current:
                return
            current["total"] = total
            current["logs"].append(mensagem)

    def worker() -> None:
        try:
            total, pasta_final, cancelado = gerar_documentos(
                caminho_planilha=caminho_planilha,
                pasta_saida_base=pasta_saida_base,
                gerar_pdf=bool(gerar_pdf),
                callback_status=callback_status,
                cancel_event=cancel_event,
            )

            with jobs_lock:
                current = jobs[job_id]
                current["total"] = total
                current["pasta_final"] = str(pasta_final)
                current["cancelado"] = cancelado
                current["status"] = "cancelado" if cancelado else "concluido"

                if cancelado:
                    current["logs"].append(
                        f"Processo cancelado. Documentos gerados antes do cancelamento: {total}"
                    )
                else:
                    current["logs"].append(f"Concluido. Total de linhas processadas: {total}")
                    current["logs"].append(f"Pasta final: {pasta_final}")

        except Exception as exc:
            with jobs_lock:
                current = jobs.get(job_id)
                if current:
                    current["status"] = "erro"
                    current["erro"] = str(exc)
                    current["logs"].append(f"Erro: {exc}")

    thread = threading.Thread(target=worker, daemon=True)
    job["thread"] = thread
    thread.start()

    return {"job_id": job_id, "mensagem": "Geracao iniciada.", "job": _serializar_job(job)}


def obter_status_job(job_id: Any) -> dict[str, Any]:
    job_id = str(job_id or "").strip()
    if not job_id:
        raise ValueError("job_id nao informado.")

    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise KeyError("Job nao encontrado.")
        return _serializar_job(job)


def cancelar_job(job_id: Any) -> dict[str, Any]:
    job_id = str(job_id or "").strip()
    if not job_id:
        raise ValueError("job_id nao informado.")

    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise KeyError("Job nao encontrado.")

        if job.get("status") != "processando":
            return _serializar_job(job)

        job["cancel_event"].set()
        job["logs"].append("Cancelamento solicitado. Aguardando finalizar o documento atual...")
        return _serializar_job(job)


def abrir_pasta(caminho: Any) -> dict[str, Any]:
    pasta = _path_obrigatorio(caminho, "Pasta")
    if not pasta.exists() or not pasta.is_dir():
        raise FileNotFoundError(f"Pasta nao encontrada: {pasta}")

    if os.name == "nt":
        os.startfile(str(pasta))  # type: ignore[attr-defined]
    elif sys.platform == "darwin":
        subprocess.Popen(["open", str(pasta)])
    else:
        subprocess.Popen(["xdg-open", str(pasta)])

    return {"path": str(pasta), "mensagem": "Pasta aberta."}


def create_flask_app():
    from flask import Flask, jsonify, request, send_from_directory

    frontend_dir = Path(__file__).resolve().parents[1] / "frontend"
    app = Flask(__name__, static_folder=None)

    @app.after_request
    def aplicar_cors(response):
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        return response

    @app.route("/api/gerador-documentos/<path:_path>", methods=["OPTIONS"])
    @app.route("/api/gerador-documentos", methods=["OPTIONS"])
    def options(_path=None):
        return ("", 204)

    def payload_json() -> dict[str, Any]:
        data = request.get_json(silent=True)
        return data if isinstance(data, dict) else {}

    def json_erro(mensagem: str, status: int = 400):
        return jsonify({"ok": False, "error": mensagem, "erro": mensagem}), status

    @app.get("/")
    @app.get("/geradordocumentos.html")
    def frontend_html():
        return send_from_directory(frontend_dir, "geradordocumentos.html")

    @app.get("/geradordocumentos.js")
    def frontend_js():
        return send_from_directory(frontend_dir, "geradordocumentos.js")

    @app.get("/api/gerador-documentos/health")
    def health():
        return jsonify({"ok": True, "servico": "gerador-documentos", **modelo_info()})

    @app.post("/api/gerador-documentos/validar")
    def api_validar():
        data = payload_json()
        try:
            info = validar_payload(data.get("planilha"), data.get("pasta_saida"), bool(data.get("gerar_pdf")))
            return jsonify({"ok": True, "mensagem": "Validacao concluida.", "info": info})
        except Exception as exc:
            return json_erro(str(exc), 400)

    @app.post("/api/gerador-documentos/gerar")
    def api_gerar():
        data = payload_json()
        try:
            result = iniciar_job(data.get("planilha"), data.get("pasta_saida"), bool(data.get("gerar_pdf")))
            return jsonify({"ok": True, **result})
        except Exception as exc:
            return json_erro(str(exc), 400)

    @app.get("/api/gerador-documentos/status/<job_id>")
    def api_status(job_id: str):
        try:
            return jsonify({"ok": True, "job": obter_status_job(job_id)})
        except KeyError as exc:
            return json_erro(str(exc), 404)
        except Exception as exc:
            return json_erro(str(exc), 400)

    @app.post("/api/gerador-documentos/cancelar/<job_id>")
    def api_cancelar(job_id: str):
        try:
            return jsonify({"ok": True, "job": cancelar_job(job_id)})
        except KeyError as exc:
            return json_erro(str(exc), 404)
        except Exception as exc:
            return json_erro(str(exc), 400)

    @app.post("/api/gerador-documentos/abrir-pasta")
    def api_abrir_pasta():
        data = payload_json()
        try:
            return jsonify({"ok": True, **abrir_pasta(data.get("caminho"))})
        except Exception as exc:
            return json_erro(str(exc), 400)

    return app


def main() -> None:
    host = "127.0.0.1"
    port = 8766
    app = create_flask_app()
    print(f"Servico Gerador de Documentos iniciado em http://{host}:{port}")
    print(f"Abra: http://{host}:{port}/geradordocumentos.html")
    print(f"Modelo esperado: {caminho_modelo_word()}")
    app.run(host=host, port=port, debug=False, threaded=True)


if __name__ == "__main__":
    main()
