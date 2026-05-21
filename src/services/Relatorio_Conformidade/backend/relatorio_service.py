from __future__ import annotations

import os
import sys
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from dateutil.relativedelta import relativedelta


def _runtime_base() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[4]  # raiz do projeto


def _resource_path(rel_path: str) -> Path:
    # PyInstaller onefile
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / rel_path
    # dev: tenta ao lado do arquivo
    here = Path(__file__).resolve().parent
    p = here / rel_path
    if p.exists():
        return p
    # fallback: src/docs
    return _runtime_base() / "src" / "docs" / rel_path


MODELO_PRECA = _resource_path("MODELO RELATORIO.docx")
MODELO_RPV = _resource_path("Conformidade  - RPV.docx")


def gerar_relatorios(excel_path: str, out_dir: str) -> dict:
    excel_path = str(excel_path or "").strip()
    out_dir = str(out_dir or "").strip()

    if not excel_path:
        return {"ok": False, "error": "excel_path vazio"}
    if not out_dir:
        return {"ok": False, "error": "out_dir vazio"}

    xls = Path(excel_path).resolve()
    out = Path(out_dir).resolve()

    if not xls.exists() or not xls.is_file():
        return {"ok": False, "error": f"Excel inválido: {xls}"}
    if not out.exists() or not out.is_dir():
        return {"ok": False, "error": f"Pasta de saída inválida: {out}"}
    if not MODELO_PRECA.exists():
        return {"ok": False, "error": f"Modelo PRECA não encontrado: {MODELO_PRECA}"}
    if not MODELO_RPV.exists():
        return {"ok": False, "error": f"Modelo RPV não encontrado: {MODELO_RPV}"}

    # subpasta automática (igual seu v3.py) :contentReference[oaicite:2]{index=2}
    pasta_final = out / f"Relatorios_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"
    pasta_final.mkdir(parents=True, exist_ok=True)

    try:
        df = pd.read_excel(str(xls))
    except Exception as e:
        return {"ok": False, "error": f"Erro ao ler Excel: {e!s}"}

    mapeamento = {
        "NUMERO_PROCESSO": "NUMERO_PROCESSO",
        "AUTOR": "AUTOR",
        "CUMPRIMENTO_SENTENCA": "CUMPRIMENTO_SENTENCA",
        "SITUACAO_PROCESSO": "SITUACAO_PROCESSO",
        "DATA_ACAO": "DATA_ACAO",
        "DATA_PERICIA": "DATA_PERICIA",
        "DATA_REALIZADA": "DATA_REALIZADA",
        "DATA_LAUDO": "DATA_LAUDO",
        "TIPO LAUDO": "TIPO LAUDO",
        "DATA_SENTENCA": "DATA_SENTENCA",
        "SENTENCA": "SENTENCA",
        "DATA_APELACAO": "DATA_APELACAO",
        "APE": "APE",
        "DATA_JULGAMENTO": "DATA_JULGAMENTO",
        "JULGA": "JULGA",
        "DATA_TRANSITO": "DATA_TRANSITO",
        "DATA_CUMPRIMENTO": "DATA_CUMPRIMENTO",
        "DATA_HOMOLOGACAO": "DATA_HOMOLOGACAO",
        "DATA_PRECA": "DATA_PRECA",
        "DATA_RPV": "DATA_RPV",
        "DATA_OFICIO": "DATA_OFICIO",
        "DATA_OR_PAGAMENTO": "DATA_OR_PAGAMENTO",
        "DATA_ENCERRAMENTO": "DATA_ENCERRAMENTO",
    }

    EVENTOS_PRECA = [
        ("DATA_ACAO", 0, 0, 0),
        ("DATA_PERICIA", 0, 2, 0),
        ("DATA_REALIZADA", 0, 1, 20),
        ("DATA_LAUDO", 0, 2, 0),
        ("DATA_SENTENCA", 0, 3, 0),
        ("DATA_APELACAO", 0, 1, 15),
        ("DATA_JULGAMENTO", 0, 2, 0),
        ("DATA_TRANSITO", 0, 1, 0),
        ("DATA_CUMPRIMENTO", 0, 1, 1),
        ("DATA_HOMOLOGACAO", 0, 3, 0),
        ("DATA_PRECA", 0, 1, 5),
        ("DATA_OFICIO", 0, 3, 0),
        ("DATA_OR_PAGAMENTO", 0, 1, 0),
        ("DATA_ENCERRAMENTO", 1, 6, 0),
    ]

    EVENTOS_RPV = [
        ("DATA_ACAO", 0, 0, 0),
        ("DATA_PERICIA", 0, 2, 0),
        ("DATA_REALIZADA", 0, 1, 20),
        ("DATA_LAUDO", 0, 2, 0),
        ("DATA_SENTENCA", 0, 3, 0),
        ("DATA_APELACAO", 0, 1, 15),
        ("DATA_JULGAMENTO", 0, 2, 0),
        ("DATA_TRANSITO", 0, 1, 0),
        ("DATA_CUMPRIMENTO", 0, 1, 1),
        ("DATA_HOMOLOGACAO", 0, 3, 0),
        ("DATA_RPV", 0, 1, 5),
        ("DATA_ENCERRAMENTO", 0, 3, 9),
    ]

    def _parse_data(valor):
        if pd.isna(valor) or valor in ("", None, "None"):
            return None
        for dayfirst in (True, False):
            try:
                return pd.to_datetime(valor, dayfirst=dayfirst, errors="raise").date()
            except Exception:
                pass
        return None

    def _fmt_dt(dt):
        return "" if dt is None else dt.strftime("%d/%m/%Y")

    def limpar_nome_arquivo(nome):
        s = str(nome or "").strip() or "SEM_NUMERO"
        for ch in ['<', '>', ':', '"', '/', '\\', '|', '?', '*']:
            s = s.replace(ch, "-")
        return s.strip() or "SEM_NUMERO"

    def norm(txt: str) -> str:
        if txt is None or (isinstance(txt, float) and pd.isna(txt)):
            return ""
        s = str(txt).strip().lower()
        s = (s.replace("á", "a").replace("à", "a").replace("â", "a").replace("ã", "a")
               .replace("é", "e").replace("ê", "e")
               .replace("í", "i")
               .replace("ó", "o").replace("ô", "o").replace("õ", "o")
               .replace("ú", "u")
               .replace("ç", "c"))
        return s

    def aplicar_marcacoes(texto: str, linha: pd.Series) -> str:
        if not texto:
            return texto
        marcacoes = {
            "LP": "( )", "LPP": "( )", "LN": "( )",
            "SENTENCA_A": "( )", "SENTENCA_I": "( )",
            "APE_A": "( )", "APE_I": "( )",
            "JULGA_A": "( )", "JULGA_I": "( )",
        }

        laudo_n = norm(linha.get("TIPO LAUDO", "") or linha.get("LAUDO", ""))
        if "positivo" in laudo_n: marcacoes["LP"] = "(X)"
        elif "parcial" in laudo_n: marcacoes["LPP"] = "(X)"
        elif "negativo" in laudo_n: marcacoes["LN"] = "(X)"

        sentenca_n = norm(linha.get("SENTENÇA", "") or linha.get("SENTENCA", ""))
        if "procedente" in sentenca_n: marcacoes["SENTENCA_A"] = "(X)"
        elif "improcedent e" in sentenca_n: marcacoes["SENTENCA_I"] = "(X)"

        apelacao_n = norm(linha.get("APELAÇÃO", "") or linha.get("APELACAO", linha.get("APE", "")))
        if "autor" in apelacao_n: marcacoes["APE_A"] = "(X)"
        elif "inss" in apelacao_n: marcacoes["APE_I"] = "(X)"

        julgamento_n = norm(linha.get("JULGAMENTO", "") or linha.get("JULGA", ""))
        if "favoravel" in julgamento_n: marcacoes["JULGA_A"] = "(X)"
        elif "desfavoravel" in julgamento_n: marcacoes["JULGA_I"] = "(X)"

        for token, repl in marcacoes.items():
            texto = texto.replace(f"({{{token}}})", repl)
        return texto

    def aplicar_fonte_calibri_light(run, cor_vermelha=False):
        run.font.name = "Calibri Light"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri Light")
        run.font.size = Pt(10.5)
        if cor_vermelha:
            run.font.color.rgb = RGBColor(255, 0, 0)

    def resolver_datas(row, sequencia_eventos):
        datas, reais = {}, []
        for i, (col, _, _, _) in enumerate(sequencia_eventos):
            dt = _parse_data(row.get(col))
            if dt is not None:
                reais.append((i, col, dt))

        if not reais:
            for col, _, _, _ in sequencia_eventos:
                datas[col] = {"valor": "", "prevista": False}
            return datas

        idx_anchor, _, dt_anchor = max(reais, key=lambda t: t[2])

        for i, (col, _, _, _) in enumerate(sequencia_eventos[: idx_anchor + 1]):
            dt = _parse_data(row.get(col))
            datas[col] = {"valor": _fmt_dt(dt) if dt else "", "prevista": False}

        cursor = dt_anchor
        for i in range(idx_anchor + 1, len(sequencia_eventos)):
            col, anos, meses, dias = sequencia_eventos[i]
            dt_real = _parse_data(row.get(col))
            if dt_real:
                datas[col] = {"valor": _fmt_dt(dt_real), "prevista": False}
                if dt_real > cursor:
                    cursor = dt_real
            else:
                cursor = cursor + relativedelta(years=anos, months=meses, days=dias)
                datas[col] = {"valor": _fmt_dt(cursor), "prevista": True}
        return datas

    def preencher_documento(doc, row, datas_resolvidas):
        # Parágrafos
        for paragraph in doc.paragraphs:
            texto_original = paragraph.text
            texto_sub = texto_original
            tem_prevista = False

            for coluna, placeholder in mapeamento.items():
                chave = f"{{{placeholder}}}"
                if chave in texto_sub:
                    if "DATA" in coluna:
                        info = datas_resolvidas.get(coluna, {"valor": "", "prevista": False})
                        texto_sub = texto_sub.replace(chave, info["valor"])
                        tem_prevista = tem_prevista or (info["prevista"] and info["valor"] != "")
                    else:
                        valor = row.get(coluna)
                        valor = "" if pd.isna(valor) or valor in ("", None, "None") else str(valor)
                        texto_sub = texto_sub.replace(chave, valor)

            texto_sub = aplicar_marcacoes(texto_sub, row)

            if texto_sub != texto_original:
                for r in paragraph.runs:
                    r.text = ""
                new_run = paragraph.add_run(texto_sub)
                aplicar_fonte_calibri_light(new_run, cor_vermelha=tem_prevista)

        # Tabelas
        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for paragraph in cell.paragraphs:
                        texto_original = paragraph.text
                        texto_sub = texto_original
                        tem_prevista = False

                        for coluna, placeholder in mapeamento.items():
                            chave = f"{{{placeholder}}}"
                            if chave in texto_sub:
                                if "DATA" in coluna:
                                    info = datas_resolvidas.get(coluna, {"valor": "", "prevista": False})
                                    texto_sub = texto_sub.replace(chave, info["valor"])
                                    tem_prevista = tem_prevista or (info["prevista"] and info["valor"] != "")
                                else:
                                    valor = row.get(coluna)
                                    valor = "" if pd.isna(valor) or valor in ("", None, "None") else str(valor)
                                    texto_sub = texto_sub.replace(chave, valor)

                        texto_sub = aplicar_marcacoes(texto_sub, row)

                        if texto_sub != texto_original:
                            for r in paragraph.runs:
                                r.text = ""
                            new_run = paragraph.add_run(texto_sub)
                            aplicar_fonte_calibri_light(new_run, cor_vermelha=tem_prevista)

        return doc

    total = len(df)
    gerados = 0
    erros = 0
    erros_lista = []

    for index, row in df.iterrows():
        numero_processo = str(row["NUMERO_PROCESSO"]) if "NUMERO_PROCESSO" in df.columns else f"_{index+1:03d}"

        try:
            # sempre gera os dois (como seu v3.py) :contentReference[oaicite:3]{index=3}
            for tipo, modelo, sequencia in (
                ("PRECA", MODELO_PRECA, EVENTOS_PRECA),
                ("RPV", MODELO_RPV, EVENTOS_RPV),
            ):
                doc = Document(str(modelo))
                datas_res = resolver_datas(row, sequencia)
                doc = preencher_documento(doc, row, datas_res)

                nome_saida = f"{tipo}_{limpar_nome_arquivo(numero_processo)}.docx"
                doc.save(str(pasta_final / nome_saida))
                gerados += 1

        except Exception as e:
            erros += 1
            erros_lista.append({"processo": str(numero_processo), "erro": str(e)})

    return {
        "ok": True,
        "out_dir": str(pasta_final),
        "total_processos": int(total),
        "arquivos_gerados": int(gerados),
        "erros": int(erros),
        "erros_lista": erros_lista[:50],
    }
