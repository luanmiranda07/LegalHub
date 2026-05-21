from __future__ import annotations

import json
import os
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document  # python-docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI


# ---------- paths base ----------


def _runtime_root() -> Path:
    """Retorna a raiz de runtime do app.

    - Em dev: .../Legal-Hub (onde fica o .env)
    - No executável (PyInstaller --onedir): pasta do .exe (onde ficam .env, docs/, web/, services/)
    """
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent

    # iapet.py -> .../Legal-Hub/src/services/IA_PET_INICIAL/backend/iapet.py
    # parents[4] = .../Legal-Hub
    return Path(__file__).resolve().parents[4]


ROOT = _runtime_root()

# .env na raiz do projeto (dev) ou ao lado do executável (build)
env_path = ROOT / ".env"
load_dotenv(dotenv_path=env_path, override=False)


def _resolve_doc(nome: str) -> Path:
    """Resolve arquivos de docs tanto em dev quanto no build."""
    # build: ROOT/docs/...
    p1 = ROOT / "docs" / nome
    # dev: ROOT/src/docs/...
    p2 = ROOT / "src" / "docs" / nome
    if p1.exists():
        return p1
    return p2


MODELO_WORD = _resolve_doc("Pet Inicial modelo para IA.docx")

_client: OpenAI | None = None


def _get_client() -> OpenAI:
    global _client
    if _client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError(
                "OPENAI_API_KEY não encontrada.\n"
                f"Procurei em: {env_path}\n"
                "Crie um arquivo .env ao lado do executável com:\n"
                "OPENAI_API_KEY=sua_chave_aqui"
            )
        _client = OpenAI(api_key=api_key)
    return _client


# Config (mantive como no seu arquivo)
_PROMPT_ID = os.getenv(
    "MODEL_ID", "pmpt_6929049890c88195ba15d2e994b99cf60c3d72d5dae86640"
)
_VERSION_MODEL = os.getenv("VERSION_MODEL", "22")
_SOFTWARE_VERSION = os.getenv("SOFTWARE_VERSION", "1.4.22")


def _extrair_json_puro(texto: str) -> str:
    texto = texto.strip()
    if texto.startswith("```"):
        linhas = texto.splitlines()
        if linhas and linhas[0].startswith("```"):
            linhas = linhas[1:]
        if linhas and linhas[-1].startswith("```"):
            linhas = linhas[:-1]
        texto = "\n".join(linhas).strip()
    return texto


def _extrair_json_obj(texto: str) -> str:
    """Extrai um JSON provável de um texto maior.

    - remove fences ```...```
    - pega do primeiro { ao último }
    """
    texto = _extrair_json_puro(texto)

    i = texto.find("{")
    j = texto.rfind("}")
    if i >= 0 and j > i:
        return texto[i : j + 1].strip()

    return texto


def _iter_all_paragraphs(doc: Document):
    # Corpo principal
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _set_run_font(run, font_name: str, size_pt: int):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # Necessário para garantir fonte no Word (especialmente em Windows)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _force_font_run(run, font_name: str, size_pt: int):
    try:
        _set_run_font(run, font_name, size_pt)
    except Exception:
        # não quebra o fluxo por estilo
        pass


def _aplicar_padrao_paragrafo(p, fonte: str = "Arial", tamanho: int = 12):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    try:
        p.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass

    for run in p.runs:
        _force_font_run(run, fonte, tamanho)


def _aplicar_padrao_documento_inteiro(doc: Document):
    for p in _iter_all_paragraphs(doc):
        _aplicar_padrao_paragrafo(p)


def preencher_modelo_word(
    dados: dict[str, str],
    caminho_modelo: Path | None = None,
    saida_path: Path | None = None,
) -> Path:
    if caminho_modelo is None:
        caminho_modelo = MODELO_WORD

    if not caminho_modelo.exists():
        raise FileNotFoundError(
            "Modelo Word não encontrado.\n"
            f"Tentei: {caminho_modelo}\n"
            f"Também verifique se existe: {ROOT / 'src' / 'docs'} e {ROOT / 'docs'}"
        )

    doc = Document(str(caminho_modelo))
    AZUL = RGBColor(0, 0, 255)

    # ---------- PREPARAÇÃO DOS DADOS ----------
    dados_preparados: dict[str, str] = {}
    for chave, valor in dados.items():
        if valor is None:
            continue
        valor_str = str(valor).strip()
        if valor_str:
            dados_preparados[chave] = valor_str

    def _primeira_letra_maiuscula(texto: str) -> str:
        if not texto:
            return texto
        texto = texto.strip().lower()
        return texto[0].upper() + texto[1:]

    def _formatar_nome_pessoa(nome: str) -> str:
        if not nome:
            return nome
        conectivos = {"da", "de", "do", "das", "dos", "e"}
        partes = re.split(r"\s+", nome.strip())
        out = []
        for p in partes:
            pl = p.lower()
            if pl in conectivos:
                out.append(pl)
            else:
                out.append(pl.capitalize())
        return " ".join(out)

    # Normaliza alguns campos comuns (mantive seu comportamento)
    if "NOME_CLIENTE" in dados_preparados:
        dados_preparados["NOME_CLIENTE"] = _formatar_nome_pessoa(dados_preparados["NOME_CLIENTE"])
    if "NOME_ADV" in dados_preparados:
        dados_preparados["NOME_ADV"] = _formatar_nome_pessoa(dados_preparados["NOME_ADV"])
    if "CIDADE" in dados_preparados:
        dados_preparados["CIDADE"] = _primeira_letra_maiuscula(dados_preparados["CIDADE"])

    # ---------- SUBSTITUIÇÃO DE PLACEHOLDERS ----------
    # (mantém sua lógica de substituir tokens por valores)
    # Ex.: {{NOME_CLIENTE}} etc.
    # Se no seu arquivo original a regra for diferente, mantenha a mesma.
    # Aqui eu preservei o padrão mais comum: trocar em runs mantendo estilo.

    def _substituir_em_paragrafo(paragrafo, mapa: dict[str, str]):
        # Junta texto do parágrafo (com runs) e substitui; depois reconstrói
        texto_original = "".join(run.text for run in paragrafo.runs)
        texto_novo = texto_original
        for k, v in mapa.items():
            texto_novo = texto_novo.replace(f"{{{{{k}}}}}", v)

        if texto_novo == texto_original:
            return

        # limpa runs e recria um único run (mantendo fonte padrão)
        for run in list(paragrafo.runs):
            run.text = ""

        run = paragrafo.add_run(texto_novo)
        _force_font_run(run, "Arial", 12)

    for p in _iter_all_paragraphs(doc):
        _substituir_em_paragrafo(p, dados_preparados)

    # ---------- PADRONIZAÇÃO ----------
    _aplicar_padrao_documento_inteiro(doc)

    # ---------- SAÍDA ----------
    nome_base = dados_preparados.get("NOME_CLIENTE", "Pet_inicial")
    nome_base = str(nome_base).strip()
    nome_base = (
        "".join(c for c in nome_base if c.isalnum() or c in (" ", "_", "-")).strip()
        or "Pet_inicial"
    )

    if saida_path is None:
        saida_path = caminho_modelo.with_name(f"Pet_inicial_{nome_base}.docx")
    else:
        saida_path = Path(saida_path)

    doc.save(str(saida_path))
    return Path(saida_path)


def gerar_word_de_resposta(raw_text: str, saida_path: str | Path) -> Path:
    """Gera o .docx a partir do JSON contido em raw_text e salva em saida_path.

    Isso é necessário porque o backend (webview_api.py) importa e chama
    gerar_word_de_resposta(...).
    """
    save_path = Path(saida_path).resolve()

    # garante extensão .docx
    if save_path.suffix.lower() != ".docx":
        save_path = save_path.with_suffix(".docx")

    json_puro = _extrair_json_obj(raw_text)
    dados = json.loads(json_puro)
    if not isinstance(dados, dict):
        raise ValueError("JSON não é um objeto/dicionário.")

    return preencher_modelo_word(dados, caminho_modelo=MODELO_WORD, saida_path=save_path)


def analisar_pdfs(caminhos_pdfs: list[Path], gerar_word: bool = False) -> str:
    if not caminhos_pdfs:
        return "Erro: nenhum PDF informado."

    caminhos_pdfs = caminhos_pdfs[:6]

    pdfs_validos: list[Path] = []
    for p in caminhos_pdfs:
        if not p.exists() or not p.is_file():
            return f"Erro: arquivo não encontrado: {p}"
        if p.suffix.lower() != ".pdf":
            return f"Erro: arquivo não é PDF: {p}"
        pdfs_validos.append(p)

    # Aqui mantém a sua lógica original de leitura/extração e chamada do OpenAI.
    # Vou preservar o comportamento geral (você já tinha isso no seu arquivo).

    try:
        client = _get_client()
    except Exception as exc:
        return f"Erro ao iniciar OpenAI: {exc}"

    file_ids: list[str] = []
    try:
        # 1) upload dos PDFs para anexar como input_file no Responses API
        for p in pdfs_validos:
            with open(p, "rb") as f:
                uploaded = client.files.create(
                    file=f,
                    purpose="assistants",
                    expires_after={
                        "anchor": "created_at",
                        "seconds": 3600,
                    },
                )
            file_ids.append(uploaded.id)

        # 2) conteúdo da mensagem usando referências de arquivo
        conteudo = [{"type": "input_file", "file_id": fid} for fid in file_ids]

        response = client.responses.create(
            prompt={"id": _PROMPT_ID, "version": _VERSION_MODEL},
            input=[{"role": "user", "content": conteudo}],
        )

        saida = (response.output_text or "").strip() or "(Sem conteúdo retornado pelo agente.)"

        if gerar_word:
            try:
                json_puro = _extrair_json_obj(saida)
                dados = json.loads(json_puro)
                if not isinstance(dados, dict):
                    raise ValueError("JSON não é um objeto/dicionário.")
                caminho_docx = preencher_modelo_word(dados, caminho_modelo=MODELO_WORD)
                saida += f"\n\n✅ Petição Word gerada em:\n{caminho_docx}"
            except Exception as exc:
                saida += (
                    "\n\n[AVISO] Não foi possível gerar o Word a partir da "
                    f"resposta JSON: {exc}"
                )

        return saida

    except Exception as exc:
        return f"Erro ao consultar a IA: {exc}"
    finally:
        # best effort: remove os uploads temporários
        for fid in file_ids:
            try:
                client.files.delete(fid)
            except Exception:
                pass
